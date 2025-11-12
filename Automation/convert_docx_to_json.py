# File: Automation/convert_docx_to_json.py
import docx
import json
import os

# --- PATHS ---
DOCX_PATH = "data game.docx"
OUTPUT_DIR = "Data"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "land_investment.json")

# --- UTILITY FUNCTION ---
def clean_text(text):
    # Membersihkan teks dari simbol mata uang, koma, dan spasi berlebihan
    return text.replace('Rp', '').replace('Coin', '').replace('.', '').replace(',', '').strip()

# --- MAIN CONVERSION LOGIC ---
def convert_land_data():
    try:
        doc = docx.Document(DOCX_PATH)
    except Exception as e:
        print(f"ERROR: Tidak dapat memuat file {DOCX_PATH}. Pastikan file ada di root folder. {e}")
        return

    land_data = []
    
    # Mencari tabel Land Investment yang biasanya ada di awal dokumen
    if not doc.tables:
        print("ERROR: Dokumen Word tidak memiliki tabel.")
        return

    # Asumsi: Tabel Land Investment adalah tabel pertama yang memiliki 4 kolom
    target_table = None
    for table in doc.tables:
        if len(table.columns) >= 4:
            target_table = table
            break
            
    if target_table is None:
        print("ERROR: Tidak menemukan tabel yang sesuai (minimal 4 kolom) untuk Land Investment.")
        return

    # Memproses baris (Melewatkan baris header)
    for i, row in enumerate(target_table.rows):
        if i < 2: # Lewatkan baris judul dan header
            continue

        cells = [clean_text(cell.text) for cell in row.cells]
        
        # Contoh: Ladang Ke-, Harga Dasar per m^2, Total Harga Investasi, Kenaikan Progresif
        if len(cells) >= 4:
            try:
                # Menangani baris GRATIS (Landang Ke-1)
                if 'GRATIS' in cells[0].upper():
                    land_data.append({
                        "id": 1,
                        "name": "Ladang Pertanian",
                        "cost": 0,
                        "base_price_m2": 0,
                        "is_free": True
                    })
                    continue

                land_id = int(cells[0].split('(')[0].strip()) # Mengambil angka dari "2 (Start Purchase)"
                
                # Menghilangkan 'Juta' dan mengubahnya menjadi angka
                total_cost_text = cells[2].replace('Juta', '000000').replace('juta', '000000').replace('Juta Coin', '')
                
                data_entry = {
                    "id": land_id,
                    "name": "Ladang Pertanian",
                    "cost": int(total_cost_text.strip()),
                    "base_price_m2": int(cells[1].strip()),
                    "progressive_increase": cells[3].strip() # Simpan sebagai teks karena mungkin deskriptif
                }
                land_data.append(data_entry)

            except ValueError as e:
                print(f"WARNING: Melewatkan baris {i+1} karena ada nilai yang tidak valid: {cells}. Error: {e}")
                
    
    # --- WRITE TO JSON ---
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        
    with open(OUTPUT_FILE, 'w') as f:
        json.dump(land_data, f, indent=2)

    print(f"✅ Sukses mengkonversi {len(land_data)} entri Land Investment ke {OUTPUT_FILE}")

if __name__ == "__main__":
    convert_land_data()

